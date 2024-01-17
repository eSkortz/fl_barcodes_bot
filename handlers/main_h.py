from aiogram import Router, F, Bot
from aiogram.types import CallbackQuery, FSInputFile, Message
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
from keyboards import start_k, only_to_main_k
from config import BOT_TOKEN, API_URL
import os
import uuid
import aiohttp
from docx import Document
from docx.shared import Pt

# Ð¡Ð¾Ð·Ð´Ð°ÐµÐ¼ ÑÐºÐ·ÐµÐ¼Ð¿Ð»ÑÑ€ Ñ€Ð¾ÑƒÑ‚ÐµÑ€Ð° Ð¸ Ð±Ð¾Ñ‚Ð°
router = Router()
bot = Bot(token=BOT_TOKEN)


# ÐžÐ¿Ñ€ÐµÐ´ÐµÐ»ÐµÐ½Ð¸Ðµ ÑÐ¾ÑÑ‚Ð¾ÑÐ½Ð¸Ð¹ Ð´Ð»Ñ finite state machine (FSM)
class Patterns(StatesGroup):
    waiting_to_id = State()


# Ð¤ÑƒÐ½ÐºÑ†Ð¸Ñ Ð¾Ð±Ñ€Ð°Ð±Ð¾Ñ‚ÐºÐ¸ Ð¾ÑˆÐ¸Ð±Ð¾Ðº
async def sth_error(message: Message, error_text: str) -> None:
    await message.delete()
    markup_inline = only_to_main_k.get()
    await message.answer(
        text=f"â›”ï¸ Ð§Ñ‚Ð¾-Ñ‚Ð¾ Ð¿Ð¾ÑˆÐ»Ð¾ Ð½Ðµ Ñ‚Ð°Ðº\n\nÐ¢ÐµÐºÑÑ‚ Ð¾ÑˆÐ¸Ð±ÐºÐ¸: {error_text}",
        reply_markup=markup_inline,
    )


# ÐžÐ±Ñ€Ð°Ð±Ð¾Ñ‚Ñ‡Ð¸Ðº Ð´Ð»Ñ Ð¾ÑÐ½Ð¾Ð²Ð½Ð¾Ð³Ð¾ Ð¼ÐµÐ½ÑŽ
@router.callback_query(F.data == "main_menu")
async def main_menu(callback: CallbackQuery) -> None:
    try:
        markup_inline = start_k.get()
        await callback.message.delete()
        await callback.message.answer(
            text="ðŸ¤– ÐŸÑ€Ð¸Ð²ÐµÑ‚, Ñ Ð±Ð¾Ñ‚ Ð´Ð»Ñ Ð¿Ð¾Ð»ÑƒÑ‡ÐµÐ½Ð¸Ñ Ð±Ð°Ñ€ÐºÐ¾Ð´Ð¾Ð² Ð² Ñ„Ð¾Ñ€Ð¼Ð°Ñ‚Ðµ docx",
            reply_markup=markup_inline,
        )
    except Exception as exception:
        await sth_error(callback.message, exception)


# ÐžÐ±Ñ€Ð°Ð±Ð¾Ñ‚Ñ‡Ð¸Ðº Ð´Ð»Ñ Ð·Ð°Ð¿Ñ€Ð¾ÑÐ° Ð³ÐµÐ½ÐµÑ€Ð°Ñ†Ð¸Ð¸ ÐºÐ¾Ð´Ð¾Ð²
@router.callback_query(F.data == "get_patterns")
async def generate(callback: CallbackQuery, state: FSMContext) -> None:
    try:
        await callback.message.delete()
        message = await callback.message.answer(
            text="ðŸ“‘ ÐŸÑ€Ð¸ÑˆÐ»Ð¸Ñ‚Ðµ Ð¼Ð½Ðµ ÐºÐ¾Ð»-Ð²Ð¾ ÐºÐ¾Ð´Ð¾Ð², ÐºÐ¾Ñ‚Ð¾Ñ€Ð¾Ðµ Ð½ÐµÐ¾Ð±Ñ…Ð¾Ð´Ð¸Ð¼Ð¾ ÑÐ³ÐµÐ½ÐµÑ€Ð¸Ñ€Ð¾Ð²Ð°Ñ‚ÑŒ"
        )
        await state.set_state(Patterns.waiting_to_id)
        await state.update_data(id_to_delete=message.message_id)
    except Exception as exception:
        await sth_error(callback.message, exception)


# ÐžÐ±Ñ€Ð°Ð±Ð¾Ñ‚Ñ‡Ð¸Ðº Ð²Ð²Ð¾Ð´Ð° ÐºÐ¾Ð»Ð¸Ñ‡ÐµÑÑ‚Ð²Ð° ÐºÐ¾Ð´Ð¾Ð²
@router.message(Patterns.waiting_to_id)
async def waiting_to_message(message: Message, state: FSMContext) -> None:
    try:
        message_text = message.text
        state_data = await state.get_data()
        id_to_delete = state_data["id_to_delete"]
        await bot.delete_message(chat_id=message.chat.id, message_id=id_to_delete)
        await message.delete()

        try:
            int(message_text)
        except Exception:
            markup_inline = only_to_main_k.get()
            await message.answer(
                text="â›”ï¸ ÐŸÑ€Ð¾Ð²ÐµÑ€ÑŒÑ‚Ðµ ÐºÐ¾Ñ€Ñ€ÐµÐºÑ‚Ð½Ð¾ÑÑ‚ÑŒ Ð²Ð²ÐµÐ´ÐµÐ½Ð½Ñ‹Ñ… Ð´Ð°Ð½Ð½Ñ‹Ñ…",
                reply_markup=markup_inline,
            )
        else:
            count = int(message_text)
            codes = generate_codes(count)
            await send_codes_to_url(url=API_URL, codes=codes)
            filename = save_codes_to_docx(codes)

            docx_file = FSInputFile(f"{filename}")
            markup_inline = only_to_main_k.get()
            await message.answer_document(
                document=docx_file, reply_markup=markup_inline
            )
            os.remove(f"{filename}")

    except Exception as exception:
        await sth_error(message, exception)


# Ð¤ÑƒÐ½ÐºÑ†Ð¸Ñ Ð³ÐµÐ½ÐµÑ€Ð°Ñ†Ð¸Ð¸ ÑÐ»ÑƒÑ‡Ð°Ð¹Ð½Ñ‹Ñ… ÐºÐ¾Ð´Ð¾Ð²
def generate_codes(count: int) -> list:
    codes = []
    for _ in range(count):
        code = str(uuid.uuid4().hex)[:8]
        codes.append(code)
    return codes


# Ð¤ÑƒÐ½ÐºÑ†Ð¸Ñ ÑÐ¾Ñ…Ñ€Ð°Ð½ÐµÐ½Ð¸Ñ ÐºÐ¾Ð´Ð¾Ð² Ð² docx Ñ„Ð°Ð¹Ð»
def save_codes_to_docx(codes: list) -> str:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(21)
    doc.styles["Normal"].font.bold = True
    for code in codes:
        doc.add_paragraph(code)
        doc.add_paragraph("------------------------------------------------------")
    filename = f"{str(uuid.uuid4())}.docx"
    doc.save(filename)
    return filename


# Ð¤ÑƒÐ½ÐºÑ†Ð¸Ñ Ð¾Ñ‚Ð¿Ñ€Ð°Ð²ÐºÐ¸ ÐºÐ¾Ð´Ð¾Ð² Ð¿Ð¾ ÑƒÐºÐ°Ð·Ð°Ð½Ð½Ð¾Ð¼Ñƒ URL
async def send_codes_to_url(codes: list, url: str) -> list:
    async with aiohttp.ClientSession() as session:
        responses = []
        for code in codes:
            body = {"code": code}
            async with session.post(url, json=body) as response:
                responses.append(await response.json())
        return responses
