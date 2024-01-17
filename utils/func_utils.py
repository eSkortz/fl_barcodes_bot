import uuid
from docx import Document
from docx.shared import Pt


def generate_codes(count: int) -> list:
    codes = []
    for _ in range(count):
        code = str(uuid.uuid4().hex)[:8]
        codes.append(code)
    return codes


def save_codes_to_docx(codes: list) -> str:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(21)
    doc.styles["Normal"].font.bold = True
    for code in codes:
        doc.add_paragraph(code)
        doc.add_paragraph("------------------------------------------------------")
    filename = f"{str(uuid.uuid4())}.docx"
    doc.save(filename)
    return filename